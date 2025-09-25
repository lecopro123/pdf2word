import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from pdf2docx import Converter
from docx.shared import Mm
from docxcompose.composer import Composer
import PyPDF2
import os
import sys

def change_label(f_path):
    dots.config(text="●")
    load_label.config(text=f"Loaded word file: {f_path}")
    return

def change_label_pdf(n_paths):
    load_label_pdf.config(text=f"{n_paths} page(s) to be converted.")
    return

def populate_data():
    convert.config(state="disabled")
    clears.config(state="disabled")

    global doc, file_path, file_paths, flag
    tmp_f_path="tempr.docx"
    count=0
    if file_paths:
        try:
            for path_to in file_paths:
                cv=Converter(path_to)
                cv.convert(tmp_f_path,start=0,end=None)
                cv.close()

                if flag:
                    doc.add_page_break()

                section = doc.sections[0]
                section.page_height = Mm(297)  # A4 height
                section.page_width = Mm(210)   # A4 width

                composer = Composer(doc)
                temp_doc=Document(tmp_f_path)
                composer.append(temp_doc)
                if count<len(file_paths)-1:
                    doc.add_page_break()
                    count+=1
        except Exception as e:
            messagebox.showerror("Error", f"Failed to trasfer PDF(s) -> Word: {e}")
        else:
            composer.save(file_path)
            messagebox.showinfo("Conversion Complete", "Your file has been successfully converted!")
        finally:
            flag=True
            if os.path.exists(tmp_f_path):
                os.remove(tmp_f_path)
                print(f"[INFO] {tmp_f_path} deleted successfully")
            else:
                print(f"[INFO] {tmp_f_path} does not exist")
            convert.config(state="normal")
            clears.config(state="normal")
    return

def clear():
    global pdf_readers, file_paths, total_pages_conv
    for widget in stack_area.winfo_children():
        widget.destroy()
    pdf_readers.clear()
    file_paths.clear()
    total_pages_conv=0
    convert.config(state="disabled")
    clears.config(state="disabled")
    # print(file_paths)
    # print(pdf_readers)
    change_label_pdf(total_pages_conv)

def delete_piece(frame,appendix):
    global total_pages_conv, pdf_readers, file_paths
    """Delete the visual piece and remove from list."""
    idx = pdf_readers.index(appendix)
    pdf_readers.pop(idx)
    file_paths.pop(idx)
    # refresh_stack_area()
    total_pages_conv=total_pages_conv-appendix['pages']
    frame.destroy()
    print(file_paths)
    change_label_pdf(total_pages_conv)
    if total_pages_conv==0:
       convert.config(state="disabled") 
       clears.config(state="disabled")
    return

def refresh_stack_area():
    """Clear and rebuild the stack area to match appendices order."""
    global pdf_readers
    for widget in stack_area.winfo_children():
        widget.destroy()
    for appendix in pdf_readers:
        create_stack_piece(appendix)
    return

def move_up(frame, appendix):
    global pdf_readers
    """Move frame and item up in the list."""
    idx = pdf_readers.index(appendix)
    if idx > 0:
        # print("entered up")
        # swap in appendices list
        pdf_readers[idx], pdf_readers[idx-1] = pdf_readers[idx-1], pdf_readers[idx]
        file_paths[idx], file_paths[idx-1] = file_paths[idx-1], file_paths[idx]
        print(file_paths)
        # move widget up
        # frame.pack_forget()
        refresh_stack_area()
    return

def move_down(frame, appendix):
    global pdf_readers
    """Move frame and item down in the list."""
    idx = pdf_readers.index(appendix)
    if idx < len(pdf_readers)-1:
        # print("entered")
        # swap in appendices list
        pdf_readers[idx], pdf_readers[idx+1] = pdf_readers[idx+1], pdf_readers[idx]
        file_paths[idx], file_paths[idx+1] = file_paths[idx+1], file_paths[idx]
        print(file_paths)
        # move widget down
        # frame.pack_forget()
        refresh_stack_area()
    return


def create_stack_piece(appendix):
    """Create a visual block (stack piece) with name and delete button."""

    frame = tk.Frame(
        stack_area,
        bd=1,
        relief="flat",
        padx=10,
        pady=5,
        bg="white",
        highlightthickness=1,
        highlightbackground="#b0b0b0",
    )
    frame.pack(fill="x", pady=6, padx=4)
    # frame = tk.Frame(stack_area, bd=2, relief="raised", padx=10, pady=5, bg="white")
    # frame.pack(fill="x", pady=5)

    def on_enter(e):
        frame.config(bg="#f5f5f5")
        lbl.config(bg="#f5f5f5")
        dot.config(bg="#f5f5f5")
        btn_frame.config(bg="#f5f5f5")

    def on_leave(e):
        frame.config(bg="white")
        lbl.config(bg="white")
        dot.config(bg="white")
        btn_frame.config(bg="white")

    frame.bind("<Enter>", on_enter)
    frame.bind("<Leave>", on_leave)

    dot = tk.Label(frame, text="●", fg="red", bg="white", font=("Arial", 10))
    dot.pack(side="left", padx=(0, 8))

    # PDF name label
    # lbl = tk.Label(frame, text=f"{appendix['name']} ({appendix['pages']} page(s))", anchor="w", bg="white")
    # lbl.pack(side="left", expand=True, fill="x")

    # File label
    lbl = tk.Label(
        frame,
        text=f"{appendix['name']} ({appendix['pages']} page(s))",
        anchor="w",
        bg="white",
        font=("Segoe UI", 10)
    )
    lbl.pack(side="left", expand=True, fill="x")

     # Right-side button frame
    btn_frame = tk.Frame(frame, bg="white")
    btn_frame.pack(side="right", padx=6)

     # Up/Down/Delete buttons with spacing
    tk.Button(
        btn_frame,
        text="▲",
        command=lambda: move_up(frame, appendix),
        width=2,
        relief="flat",
        bg="white"
    ).pack(side="left", padx=4)

    tk.Button(
        btn_frame,
        text="▼",
        command=lambda: move_down(frame, appendix),
        width=2,
        relief="flat",
        bg="white"
    ).pack(side="left", padx=4)

    tk.Button(
        btn_frame,
        text="❌",
        fg="red",
        command=lambda: delete_piece(frame, appendix),
        width=2,
        relief="flat",
        bg="white"
    ).pack(side="left", padx=4)

    # tk.Button(frame, text="▲", command=lambda: move_up(frame, appendix), width=2).pack(side="right")
    # tk.Button(frame, text="▼", command=lambda: move_down(frame, appendix), width=2).pack(side="right")
    # # Red cross button
    # btn = tk.Button(frame, text="❌", fg="red", command=lambda: delete_piece(frame,appendix))
    # btn.pack(side="right")
    convert.config(state="normal")
    clears.config(state="normal")
    return

def open_pdf_files():
    """Let user choose one or more PDF files and load them."""
    global  pdf_readers, file_paths, doc, total_pages_conv #, pdf_file_objs, pdf_readers,
    file_paths.extend(list(filedialog.askopenfilenames(
        title="Open PDF Files",
        filetypes=[("PDF Files", "*.pdf")]
    )))
    print(file_paths)
    pdf_readers = []
    total_pages_conv=0
    #pdf_file_objs=[]
    if file_paths:
        try:
            for i,path in enumerate(file_paths):
                # f = open(path, "rb")
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    pages = len(reader.pages)
                    total_pages_conv+=pages
                    name_of_pdf = path.split("/")[-1]
                    appendix = {"name": name_of_pdf, "path":path, "pages": pages, "id": i}
                    pdf_readers.append(appendix)
                    # create_stack_piece(appendix)
            # messagebox.showinfo("Loaded", f"Loaded {len(file_paths)} PDF file(s).")
            # change_label_pdf(len(file_paths))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load PDF: {e}")
        else:
            refresh_stack_area()
            change_label_pdf(total_pages_conv)
    return
            

def open_doc_file():
    """Let user choose a Word file and load it."""
    global doc, file_path, flag
    file_path = filedialog.askopenfilename(
        title="Open Word Document",
        filetypes=[("Word Files", "*.docx")]
    )
    if file_path:
        try:
            flag=True
            print(f"[INFO] {file_path}")
            doc = Document(file_path)
            # change_label(file_path)
            # pdf_but.config(state="normal")
            # messagebox.showinfo("Loaded", f"Loaded: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open the word file: {e}")
        else:
            change_label(file_path)
            pdf_but.config(state="normal")
    return

def create_doc_file():
    # Ask user for file path (like Save As)
    global doc, file_path, flag
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="Create New Word File"
    )
    if file_path:  # if user didn't cancel
        try:
            flag=False
            print(f"[INFO] {file_path}")
            doc = Document()
            doc.save(file_path) 
            #doc.save(file_path)
            # messagebox.showinfo("Created", f"New file created:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create word file: {e}")
        else:
            change_label(file_path)
            pdf_but.config(state="normal")
    return


# def save_ert():
#     """Append user input as an ert and save new file."""
#     if not file_path:
#         messagebox.showerror("Error", "No file loaded!")
#         return
    
#     ert_text = text_area.get("1.0", tk.END).strip()
#     if not ert_text:
#         messagebox.showwarning("Empty", "ert text is empty!")
#         return
    
#     # Add appendix at the end
#     doc.add_page_break()
#     doc.add_heading("ERT", level=1)
#     doc.add_paragraph(ert_text)

#     # Save as new file
#     save_path = filedialog.asksaveasfilename(
#         title="Save Word File",
#         defaultextension=".docx",
#         filetypes=[("Word Files", "*.docx")]
#     )
#     if save_path:
#         doc.save(save_path)
#         messagebox.showinfo("Saved", f"Saved to: {save_path}")

# ---------------- GUI ----------------
# states="disabled"
root = tk.Tk()
root.title("PDF to DOCX converter")
root.geometry("600x400")

root_frame = tk.Frame(root)
root_frame.pack(pady=10)

tk.Button(root_frame, text="Open Word File", command=open_doc_file).pack(side="left",padx=5)
tk.Button(root_frame, text="Create Word File", command=create_doc_file).pack(side="left",padx=5)

pdf_but=tk.Button(root_frame, text="Select PDF File(s)", command=open_pdf_files, state="disabled")
pdf_but.pack(side="right",padx=5)

status_frame = tk.Frame(root)
status_frame.pack(pady=5)

dots=tk.Label(status_frame, text="", fg="blue", font=("Arial", 12))
dots.pack(side="left")
load_label = tk.Label(status_frame, text="No active word file", font=("Segoe UI", 10))
load_label.pack(side="left", padx=(4,0))

load_label_pdf=tk.Label(root, text=f"No active pdf file(s)")
load_label_pdf.pack()

###---------Scrollable area-------------
stack_box = tk.LabelFrame(root, text="Selected PDF Appendices", padx=5, pady=5)
stack_box.pack(fill="x", padx=10, pady=10)

container = tk.Frame(stack_box, height=200)  # fixed height box
container.pack(fill="x", expand=False)

container.pack_propagate(False)

canvas = tk.Canvas(container, bg="#f9f9f9", highlightthickness=0)
canvas.pack(side="left", fill="both", expand=True)

scrollbar = tk.Scrollbar(container, orient="vertical")
scrollbar.pack(side="right", fill="y")

canvas.configure(yscrollcommand=scrollbar.set)

stack_area = tk.Frame(canvas, bg="#f9f9f9")
canvas_window=canvas.create_window((0, 0), window=stack_area, anchor="nw", width=560)  # width to fit nicely

def update_scrollregion():
    """Update the scrollable area size."""
    canvas.update_idletasks()
    bbox = canvas.bbox("all")
    if bbox:
        x1, y1, x2, y2 = bbox
        # get visible height of the canvas (the box height)
        visible_height = canvas.winfo_height()
        # clamp the bottom so scrollregion is never smaller than viewport
        if (y2 - y1) < visible_height:
            y2 = y1 + visible_height
        canvas.config(scrollregion=(x1, y1, x2, y2))

def resize_canvas(event):
    canvas.itemconfig(canvas_window, width=event.width)

stack_area.bind("<Configure>", lambda e: update_scrollregion())
canvas.bind("<Configure>", resize_canvas)

# Cross-platform scroll bindings
def _on_mousewheel(event):
    if sys.platform == "darwin":  # macOS
        canvas.yview_scroll(-1 * event.delta, "units")
    else:  # Windows & others
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

canvas.bind_all("<MouseWheel>", _on_mousewheel)
canvas.bind_all("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))  # Linux up
canvas.bind_all("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))   # Linux down
# --- End scrollable area ---
btn_frame = tk.Frame(root)
btn_frame.pack(pady=10)

clears = tk.Button(btn_frame, text="Clear", width=8, command=clear, state="disabled")
clears.pack(side="left",padx=10)

convert = tk.Button(btn_frame, text="Convert", width=8, command=populate_data, state="disabled")
convert.pack(side="left",padx=10)


file_path = None
file_paths = []
doc = None
flag=True

root.mainloop()
